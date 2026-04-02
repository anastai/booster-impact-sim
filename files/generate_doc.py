"""Generate Word documentation for the Booster Impact Simulation — Velocity Frame version."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helpers ─────────────────────────────────────────────────────────────────
def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(text)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9.5)
    return p

def shade_cell(cell, fill_hex):
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'),  fill_hex)
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:val'),   'clear')
    cell._tc.get_or_add_tcPr().append(shd)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.size = Pt(10)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(c, 'D9E1F2')
    for r, row_data in enumerate(rows):
        row = table.rows[r + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if r % 2 == 1:
                shade_cell(cell, 'F2F2F2')
    if col_widths:
        for row in table.rows:
            for j, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[j])
    return table

# ════════════════════════════════════════════════════════════════════════════
# COVER
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Satellite Launcher — Booster Impact Tracker')
r.bold = True; r.font.size = Pt(26); r.font.color.rgb = RGBColor(30, 80, 160)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Technical Documentation — Velocity Frame Edition')
r2.font.size = Pt(14); r2.font.color.rgb = RGBColor(100, 100, 100)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run(f'Version 2.0 — {datetime.date(2026, 4, 2).strftime("%B %d, %Y")}')
r3.font.size = Pt(11); r3.font.color.rgb = RGBColor(130, 130, 130)

doc.add_paragraph()
doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '1. Table of Contents', level=1)
toc = [
    '2. Overview',
    '3. Purpose and Scope',
    '4. Coordinate Systems',
    '   4.1  Inertial (Earth-Fixed) Frame',
    '   4.2  Velocity Frame — Definition',
    '   4.3  Frame Transformation (Inertial ↔ Velocity)',
    '5. Physics and Mathematical Model',
    '   5.1  Equations of Motion in the Velocity Frame',
    '   5.2  Gravity Turn Dynamics',
    '   5.3  Drag Force (velocity frame)',
    '   5.4  Lift Force (velocity frame)',
    '   5.5  Thrust Decomposition',
    '   5.6  Acceleration Commands',
    '   5.7  Atmosphere Models',
    '   5.8  Gravity Model',
    '6. Input Parameters',
    '   6.1  Booster Parameters',
    '   6.2  Motor Parameters',
    '   6.3  Aerodynamics & Atmosphere',
    '   6.4  Trajectory & Launch Angles',
    '   6.5  Velocity-Frame Guidance (New)',
    '7. Simulation Algorithm',
    '8. Output Metrics',
    '9. Visualization Charts',
    '10. Default Configuration',
    '11. Usage Guide',
    '12. Limitations and Assumptions',
    '13. Version History',
]
for item in toc:
    p = doc.add_paragraph(item)
    p.paragraph_format.left_indent = Cm(0.8 if item.startswith('   ') else 0)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# 2. OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '2. Overview', level=1)
add_body(doc,
    'The Booster Impact Tracker (v2.0) is a browser-based interactive simulation that '
    'predicts where a rocket booster will land after a suborbital launch. Version 2.0 '
    'replaces the original inertial-frame Cartesian integration with a full velocity-frame '
    'formulation, adding gravity-turn dynamics, aerodynamic lift, and 3-D heading capability.')
add_body(doc,
    'The tool is implemented as a single self-contained HTML file (booster-impact.html) '
    'and runs entirely in any modern web browser. Interactive sliders control all physical '
    'parameters; six chart views display trajectory, altitude, velocity, forces, '
    'acceleration, and flight-path angles after each run.')

# ════════════════════════════════════════════════════════════════════════════
# 3. PURPOSE AND SCOPE
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '3. Purpose and Scope', level=1)
add_body(doc,
    'The tool is intended for educational and preliminary mission-planning purposes. '
    'Primary use cases:')
for item in [
    'Estimating booster impact zones for range-safety analysis',
    'Studying gravity-turn pitch-over and its effect on impact range',
    'Comparing guided vs. unguided (ballistic) trajectories via velocity-frame acceleration commands',
    'Exploring 3-D heading changes caused by lateral (yaw-channel) acceleration commands',
    'Classroom demonstration of velocity-frame mechanics, lift/drag forces, and flight-path angle evolution',
]:
    add_bullet(doc, item)

# ════════════════════════════════════════════════════════════════════════════
# 4. COORDINATE SYSTEMS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '4. Coordinate Systems', level=1)

add_heading(doc, '4.1 Inertial (Earth-Fixed) Frame', level=2)
add_body(doc,
    'A flat-Earth, right-handed Cartesian frame fixed to the launch site:')
rows = [
    ['x', 'North  (m)',    'Downrange direction, aligned with initial heading'],
    ['y', 'East   (m)',    'Cross-range direction'],
    ['h', 'Altitude (m)', 'Vertical direction, positive upward'],
]
add_table(doc, ['Axis', 'Direction', 'Description'], rows, col_widths=[2, 3.5, 11.5])
doc.add_paragraph()

add_heading(doc, '4.2 Velocity Frame — Definition', level=2)
add_body(doc,
    'The velocity frame is centred on the vehicle and oriented with one axis along the '
    'instantaneous velocity vector. It is defined by two angles:')
rows = [
    ['γV (gammaV)', 'Flight-path angle',
     'Angle of the velocity vector above the local horizontal plane. '
     'Positive = climbing. Evolves continuously during flight.'],
    ['γH (gammaH)', 'Heading azimuth',
     'Angle of the velocity vector projected onto the horizontal plane, '
     'measured from North (0°) clockwise. Changes when lateral commands are applied.'],
]
add_table(doc, ['Symbol', 'Name', 'Description'], rows, col_widths=[3, 3.5, 10.5])
doc.add_paragraph()

add_body(doc, 'Three orthogonal unit vectors define the velocity frame:')
rows = [
    ['ê_t',  'Tangential',      'Along the velocity vector — positive in direction of motion'],
    ['ê_nV', 'Normal-vertical', 'Perpendicular to v in the pitch plane, pointing "above" the velocity vector'],
    ['ê_nH', 'Normal-lateral',  'Completing the right-handed triad (yaw / cross-track direction)'],
]
add_table(doc, ['Unit vector', 'Name', 'Description'], rows, col_widths=[3, 3.5, 10.5])
doc.add_paragraph()

add_heading(doc, '4.3 Frame Transformation (Inertial ↔ Velocity)', level=2)
add_body(doc,
    'The rotation from inertial to velocity frame is a sequence of two rotations: '
    'first by −γH about the vertical axis (heading), then by −γV about the lateral axis (pitch). '
    'The velocity vector in the inertial frame is:')
add_code(doc, 'v_inertial = v · [ cos(γV)·cos(γH),  cos(γV)·sin(γH),  sin(γV) ]')
add_body(doc, 'Position is integrated in the inertial frame using:')
add_code(doc, 'dh/dt = v · sin(γV)')
add_code(doc, 'dx/dt = v · cos(γV) · cos(γH)   (north)')
add_code(doc, 'dy/dt = v · cos(γV) · sin(γH)   (east)')
add_body(doc,
    'Forces computed in the velocity frame are transformed back to inertial components '
    'for the scalar speed and angle-rate equations (Section 5.1).')

# ════════════════════════════════════════════════════════════════════════════
# 5. PHYSICS AND MATHEMATICAL MODEL
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '5. Physics and Mathematical Model', level=1)

add_heading(doc, '5.1 Equations of Motion in the Velocity Frame', level=2)
add_body(doc,
    'The state vector is (v, γV, γH, h, x, y). '
    'Three scalar ODEs govern the velocity-frame dynamics:')

add_body(doc, 'Tangential (speed along ê_t):')
add_code(doc, 'm · (dv/dt) = T_t − D − m·g·sin(γV) + m·a_t_cmd')

add_body(doc, 'Pitch-normal (flight-path angle rate, ê_nV):')
add_code(doc, 'm·v · (dγV/dt) = T_nV + L − m·g·cos(γV) + m·a_nV_cmd')

add_body(doc, 'Yaw-normal (heading rate, ê_nH):')
add_code(doc, 'm·v·cos(γV) · (dγH/dt) = T_nH + m·a_nH_cmd')

add_body(doc,
    'where T_t, T_nV, T_nH are thrust components in the velocity frame (Section 5.5), '
    'D is drag (Section 5.3), L is lift (Section 5.4), g is altitude-dependent gravity '
    '(Section 5.8), and a_*_cmd are external acceleration commands (Section 5.6).')

add_heading(doc, '5.2 Gravity Turn Dynamics', level=2)
add_body(doc,
    'A gravity turn is achieved when thrust is aligned with the velocity vector '
    '(T_nV = T_nH = 0, T_t = F). The pitch equation reduces to:')
add_code(doc, 'dγV/dt = L/(m·v) − g·cos(γV)/v')
add_body(doc,
    'Without lift (CL = 0) this becomes the pure gravity-turn equation:')
add_code(doc, 'dγV/dt = −g · cos(γV) / v')
add_body(doc,
    'The term −g·cos(γV)/v is always negative for 0 < γV < 90°, '
    'so the flight-path angle decreases continuously after liftoff — the rocket '
    'naturally "tips over" toward horizontal due to gravity without any active steering. '
    'The rate of pitch-over is fastest at low speeds (large denominator effect is smaller) '
    'and slows as the vehicle accelerates.')
add_body(doc,
    'When gravity turn is OFF, the body is held at the fixed launch angle γV₀. '
    'The angle of attack α = γV₀ − γV grows as γV decreases, generating a normal '
    'thrust component T_nV = F·sin(α) that partially opposes the gravity-induced pitch-over.')

add_heading(doc, '5.3 Drag Force (velocity frame)', level=2)
add_body(doc,
    'Drag acts entirely along −ê_t (opposing motion). Its magnitude is:')
add_code(doc, 'D = ½ · ρ(h) · v² · CD · Aref')
add_body(doc,
    'During powered ascent CD = user-set drag coefficient. '
    'After burnout CD switches to CD_fall (higher tumbling value). '
    'Aref = π·(d/2)² is the cross-sectional reference area.')

add_heading(doc, '5.4 Lift Force (velocity frame)', level=2)
add_body(doc,
    'Lift acts along +ê_nV (perpendicular to velocity in the pitch plane, "upward from the '
    'velocity vector"). Its magnitude is:')
add_code(doc, 'L = ½ · ρ(h) · v² · CL · Aref')
add_body(doc,
    'CL is the user-controlled lift coefficient (default 0). Lift contributes directly to '
    'the flight-path angle rate (dγV/dt), slowing the gravity-turn pitch-over when CL > 0.')

add_heading(doc, '5.5 Thrust Decomposition', level=2)
add_body(doc, 'Two modes are available:')
rows = [
    ['Gravity turn ON\n(checkbox checked)',
     'Thrust ∥ ê_t\nT_t = F,  T_nV = 0,  T_nH = 0',
     'Ideal gravity turn. No angle-of-attack losses. Engine gimbal always aligned with velocity.'],
    ['Gravity turn OFF\n(checkbox unchecked)',
     'Body at fixed (γV₀, γH₀).\nα = γV₀ − γV (angle of attack)\nT_t  = F·cos(α)\nT_nV = F·sin(α)\nT_nH = 0',
     'Original fixed-angle model. Thrust steers vehicle back toward initial launch angle as γV decreases.'],
]
add_table(doc, ['Mode', 'Thrust components', 'Description'], rows, col_widths=[4, 5, 8])
doc.add_paragraph()

add_heading(doc, '5.6 Acceleration Commands', level=2)
add_body(doc,
    'Three constant commanded accelerations can be applied throughout the flight '
    'to simulate guidance corrections or lateral manoeuvres:')
rows = [
    ['a_t_cmd',   'Tangential ê_t',   'm/s²', '−10 … +10', '0',
     'Augments speed change rate (e.g. simulates throttling or drag braking)'],
    ['a_nV_cmd',  'Pitch-normal ê_nV','m/s²', '−10 … +10', '0',
     'Directly modifies dγV/dt; positive pulls flight path upward'],
    ['a_nH_cmd',  'Yaw-normal ê_nH',  'm/s²', '−5 … +5',   '0',
     'Modifies dγH/dt; creates lateral divert and changes downrange heading'],
]
add_table(doc,
    ['Parameter', 'Axis', 'Unit', 'Range', 'Default', 'Effect'],
    rows, col_widths=[2.8, 3.2, 1.5, 2.5, 1.8, 5.2])
doc.add_paragraph()

add_body(doc,
    'In the equations of motion the commanded accelerations appear as additional '
    'body-force terms: a_t_cmd enters the dv/dt equation directly; '
    'a_nV_cmd contributes a_nV_cmd/v to dγV/dt; '
    'a_nH_cmd contributes a_nH_cmd / (v·cos(γV)) to dγH/dt.')

add_heading(doc, '5.7 Atmosphere Models', level=2)
rows = [
    ['ISA Standard (default)', 'International Standard Atmosphere: troposphere lapse 6.5 K/km to 11 km, '
     'isothermal stratosphere to 20 km, upper stratosphere to 86 km. ρ → 0 above 86 km.'],
    ['Exponential', 'Scale-height model: ρ = 1.225 · exp(−h / 8500) kg/m³'],
    ['Vacuum', 'ρ = 0 everywhere. No drag or lift. Gives pure ballistic upper bound.'],
]
add_table(doc, ['Model', 'Description'], rows, col_widths=[4, 13])
doc.add_paragraph()

add_heading(doc, '5.8 Gravity Model', level=2)
add_code(doc, 'g(h) = g₀ · (Rₑ / (Rₑ + h))²')
add_body(doc,
    'where g₀ = 9.80665 m/s² and Rₑ = 6 371 000 m. '
    'In the velocity frame this produces both tangential deceleration (−g·sin γV) '
    'and pitch-down rotation (−g·cos γV / v).')

# ════════════════════════════════════════════════════════════════════════════
# 6. INPUT PARAMETERS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '6. Input Parameters', level=1)

add_heading(doc, '6.1 Booster Parameters', level=2)
rows = [
    ['Payload mass',   'mPay',  '50–2000',  '50',  '300',  'kg',
     'Mass of payload (upper stage + satellite)'],
    ['Propellant mass','mProp', '500–30000','500', '5000', 'kg',
     'First-stage propellant loaded'],
    ['Structural mass','mStr',  '200–5000', '100', '800',  'kg',
     'Dry structural mass of first stage'],
]
add_table(doc, ['Parameter','ID','Range','Step','Default','Unit','Description'], rows,
    col_widths=[3,1.8,2.2,1.5,2,1.2,5.3])
doc.add_paragraph()

add_heading(doc, '6.2 Motor Parameters', level=2)
rows = [
    ['Specific impulse','isp',     '200–420', '5',  '260','s',  'Engine efficiency'],
    ['Thrust',          'thrust',  '20–800',  '10', '120','kN', 'Constant sea-level thrust'],
    ['Burn duration',   'burnTime','10–300',  '5',  '60', 's',  'Max engine firing time'],
]
add_table(doc, ['Parameter','ID','Range','Step','Default','Unit','Description'], rows,
    col_widths=[3,1.8,2.2,1.5,2,1.2,5.3])
doc.add_paragraph()

add_heading(doc, '6.3 Aerodynamics & Atmosphere', level=2)
rows = [
    ['Drag coeff. CD',   'cd',     '0.10–1.20','0.05','0.40','—', 'Axial drag (powered ascent)'],
    ['Lift coeff. CL',   'cl',     '0.00–1.50','0.05','0.00','—', 'Lift coefficient; 0 = no lift'],
    ['Ref. diameter',    'diam',   '0.3–4.0',  '0.1', '1.2', 'm', 'Body diameter for Aref'],
    ['Post-burnout CD',  'cdFall', '0.5–3.0',  '0.1', '1.2', '—', 'Tumbling drag after burnout'],
    ['Atmosphere model', 'atm',    'ISA/Exp/Vacuum','—','ISA','—', 'ρ(h) model'],
]
add_table(doc, ['Parameter','ID','Range','Step','Default','Unit','Description'], rows,
    col_widths=[3.2,1.6,2.8,1.5,1.8,1.2,4.9])
doc.add_paragraph()

add_heading(doc, '6.4 Trajectory & Launch Angles', level=2)
rows = [
    ['Launch angle γV₀', 'angle',   '30–90',  '1', '75', '°', 'Initial flight-path angle (elevation)'],
    ['Launch azimuth γH₀','azimuth','0–359',   '1', '90', '°', 'Initial heading (0°=N, 90°=E)'],
    ['Launch latitude',  'lat',     '−45–45',  '1', '28', '°', 'Informational; does not affect physics'],
]
add_table(doc, ['Parameter','ID','Range','Step','Default','Unit','Description'], rows,
    col_widths=[3.2,1.8,2,1.5,2,1.2,5.3])
doc.add_paragraph()

add_heading(doc, '6.5 Velocity-Frame Guidance  (New in v2.0)', level=2)
rows = [
    ['Gravity turn',       'gravTurn','checked / unchecked','—','Checked','—',
     'ON: thrust ∥ velocity. OFF: thrust at fixed body angle γV₀.'],
    ['a_t command',        'aCmdT',  '−10 … +10','0.5','0','m/s²',
     'Commanded tangential acceleration along ê_t'],
    ['a_nV command',       'aCmdNV', '−10 … +10','0.5','0','m/s²',
     'Commanded pitch-normal acceleration along ê_nV'],
    ['a_nH command',       'aCmdNH', '−5 … +5',  '0.25','0','m/s²',
     'Commanded yaw-normal acceleration along ê_nH'],
]
add_table(doc, ['Parameter','ID','Range','Step','Default','Unit','Description'], rows,
    col_widths=[3,1.8,2.5,1.5,1.5,1.5,5.2])
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 7. SIMULATION ALGORITHM
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '7. Simulation Algorithm', level=1)
add_body(doc,
    'The run() function executes a forward-Euler integration loop with Δt = 0.2 s '
    '(up to 60 000 steps ≈ 3.3 h max). Each iteration:')
steps = [
    '1. Compute ρ(h) from selected atmosphere model.',
    '2. Compute dynamic pressure q = ½ρv².',
    '3. Compute drag D = q · CD · Aref  and  lift L = q · CL · Aref.',
    '4. Compute altitude-dependent gravity g(h).',
    '5. If engine on and propellant remains: apply thrust F, consume Δm = ṁ·Δt from propellant and total mass. Shut off engine when propellant exhausted or burn-time limit reached; record burnout state (h, t, γV).',
    '6. Decompose thrust into velocity-frame components T_t, T_nV, T_nH (gravity-turn or fixed-angle mode).',
    '7. Evaluate the three velocity-frame ODEs: dv/dt, dγV/dt, dγH/dt (with command augmentations).',
    '8. Integrate: v += dv/dt·Δt;  γV += dγV/dt·Δt;  γH += dγH/dt·Δt.  Clamp v ≥ 0, |γV| < 90°.',
    '9. Update inertial position: h, x, y using velocity components.',
    '10. Record every 5th sample to output arrays for charting.',
    '11. Terminate when h < 0 (impact) or loop limit reached.',
]
for s in steps:
    add_bullet(doc, s)

# ════════════════════════════════════════════════════════════════════════════
# 8. OUTPUT METRICS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '8. Output Metrics', level=1)
rows = [
    ['Max altitude',       'km',   'Peak altitude during flight'],
    ['Max velocity',       'm/s',  'Maximum speed (typically near end of burn)'],
    ['Burnout altitude',   'km',   'Altitude when engine stops'],
    ['Impact range',       'km',   '3-D total downrange: √(x²+y²) at h=0'],
    ['γV at burnout',      '°',    'Flight-path angle at engine cutoff  (NEW)'],
    ['Max lift',           'kN',   'Peak aerodynamic lift force  (NEW)'],
]
add_table(doc, ['Metric','Unit','Description'], rows, col_widths=[4,2,11])
doc.add_paragraph()
add_body(doc,
    'Impact range now uses the 3-D horizontal distance √(x²+y²), '
    'so lateral heading manoeuvres driven by a_nH commands are reflected correctly.')

# ════════════════════════════════════════════════════════════════════════════
# 9. VISUALIZATION CHARTS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '9. Visualization Charts', level=1)
rows = [
    ['Trajectory',     'Scatter (x–y)',     'Downrange (km)','Altitude (km)',
     'Full flight arc from launch to impact'],
    ['Altitude',       'Line vs time',      'Time (s)','Altitude (km)',
     'Altitude profile — ascent, coast, descent'],
    ['Velocity',       'Line vs time',      'Time (s)','Velocity (m/s)',
     'Total speed — acceleration during burn, deceleration from drag'],
    ['Forces',         'Multi-line vs time','Time (s)','Force (kN)',
     'Thrust (solid), Drag (dashed), Lift (dotted) — NEW: lift series added'],
    ['Acceleration',   'Line vs time',      'Time (s)','Accel (g)',
     'Total inertial acceleration magnitude = √(dv²+(v·dγV)²+(v·cos(γV)·dγH)²)/g₀'],
    ['Flight angles γ','Multi-line vs time','Time (s)','Angle (°)',
     'γV (flight-path angle) and γH (heading azimuth) vs time — NEW'],
]
add_table(doc,
    ['Tab','Chart Type','X-Axis','Y-Axis','What it Shows'],
    rows, col_widths=[3,3,3,3,5])
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 10. DEFAULT CONFIGURATION
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '10. Default Configuration', level=1)
rows = [
    ['Payload mass',       '300 kg'],
    ['Propellant mass',    '5 000 kg'],
    ['Structural mass',    '800 kg'],
    ['Specific impulse',   '260 s'],
    ['Thrust',             '120 kN'],
    ['Burn duration',      '60 s'],
    ['Drag coeff. CD',     '0.40'],
    ['Lift coeff. CL',     '0.00'],
    ['Reference diameter', '1.2 m'],
    ['Post-burnout CD',    '1.2 (tumbling)'],
    ['Atmosphere model',   'ISA Standard'],
    ['Launch angle γV₀',   '75°'],
    ['Launch azimuth γH₀', '90° (East)'],
    ['Launch latitude',    '28°'],
    ['Gravity turn',       'ON'],
    ['a_t command',        '0 m/s²'],
    ['a_nV command',       '0 m/s²'],
    ['a_nH command',       '0 m/s²'],
]
add_table(doc, ['Parameter','Default Value'], rows, col_widths=[7, 10])
doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# 11. USAGE GUIDE
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '11. Usage Guide', level=1)
steps = [
    'Open booster-impact.html in a modern web browser. The simulation runs automatically with default values.',
    'Set booster and motor parameters (mass, Isp, thrust, burn time) in Card 1.',
    'Set aerodynamic coefficients (CD, CL, diameter, tumbling CD) and atmosphere model in Card 2.',
    'Set launch angles (γV₀ elevation, γH₀ azimuth) in Card 2.',
    'Choose Gravity Turn mode (Card 3). Leave checked for realistic trajectory; uncheck to hold fixed body angle.',
    'Optionally set velocity-frame acceleration commands (a_t, a_nV, a_nH) to simulate guidance manoeuvres.',
    'Press "Run simulation". The six metric cards update and the impact detail strip appears.',
    'Click chart tabs to explore Trajectory, Altitude, Velocity, Forces, Acceleration, or Flight Angles.',
    'Tip: Enable CL > 0 and observe the "Flight angles γ" chart — γV decreases more slowly when lift partially counteracts gravity turn.',
    'Tip: Set a_nH ≠ 0 and observe γH drifting in the "Flight angles γ" chart; impact range uses 3-D horizontal distance.',
]
for i, s in enumerate(steps, 1):
    add_bullet(doc, f'Step {i}: {s}')

# ════════════════════════════════════════════════════════════════════════════
# 12. LIMITATIONS AND ASSUMPTIONS
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '12. Limitations and Assumptions', level=1)
for lim in [
    'Flat-Earth geometry: Earth curvature is not modelled. Results degrade beyond ~500 km downrange.',
    'Forward-Euler integration (Δt = 0.2 s): first-order accuracy. Sufficient for qualitative studies; not suitable for precision trajectory design.',
    'Constant commanded accelerations: a_t/a_nV/a_nH are applied as fixed values throughout the flight. A real guidance law would compute these dynamically.',
    'No wind: atmospheric wind is not included.',
    'Instantaneous angle of attack: in fixed-angle mode α = γV₀ − γV is applied without aerodynamic delay or structural load limits.',
    'No parachute or recovery system after burnout.',
    'Single stage only: multi-stage vehicles require separate runs per stage.',
    'Launch latitude is informational: Coriolis and Earth-rotation velocity contribution are not applied.',
    'No terrain model: impact defined at h = 0 (sea level).',
    'Lift coefficient CL is constant (no variation with angle of attack or Mach number).',
]:
    add_bullet(doc, lim)

# ════════════════════════════════════════════════════════════════════════════
# 13. VERSION HISTORY
# ════════════════════════════════════════════════════════════════════════════
add_heading(doc, '13. Version History', level=1)
rows = [
    ['1.0', 'March 26 2026',
     'Initial release. Inertial-frame Cartesian integration (vV, vH). '
     'Fixed launch angle. ISA/Exp/Vacuum atmosphere. 5 chart tabs.'],
    ['2.0', 'April 2 2026',
     'Velocity-frame reformulation (v, γV, γH). Gravity turn toggle. '
     'Aerodynamic lift (CL). Three velocity-frame acceleration commands (a_t, a_nV, a_nH). '
     '3-D heading and impact range. New γV-at-burnout and Max-lift metrics. '
     'New "Flight angles γ" chart tab. Forces chart updated with lift series.'],
]
add_table(doc, ['Version','Date','Changes'], rows, col_widths=[2, 3.5, 11.5])
doc.add_paragraph()

# ── footer ───────────────────────────────────────────────────────────────────
doc.add_page_break()
fp = doc.add_paragraph()
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = fp.add_run(
    'Satellite Launcher — Booster Impact Tracker  |  Technical Documentation v2.0  |  April 2026')
fr.font.size = Pt(9)
fr.font.color.rgb = RGBColor(150, 150, 150)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r'C:\Users\Anastasia\OneDrive\Desktop\Satellite_sim\files\Booster_Impact_Simulation_Documentation.docx'
doc.save(out)
print(f'Saved: {out}')
