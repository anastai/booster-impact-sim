"""
Regenerate Figures 1, 3, 4 with v3.0 direct-collocation optimizer,
replace images in the docx, update captions.
"""
import os, sys
os.chdir(os.path.dirname(os.path.abspath(__file__)))

# ── Use non-interactive backend so plt.show() is a no-op ─────────────────────
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

from sim.physics import Params
from sim.opt_trajectory import optimize_trajectory, plot_optimised
from grid_run import run_grid, plot_grid

# ── Shared vehicle params (same as notebook §16) ─────────────────────────────
BASE = Params(
    m_pay=10, m_prop=100, m_str=50, isp=260, thrust_kn=31, burn_max=20,
    cd=0.4, cd_fall=0.4, cl=0.0, cd_ctrl=0.3, diam=0.5,
    launch_angle=45, launch_azimuth=45,
    x_target=20, y_target=20, z_target=0, a_lat_max=20.0,
)
OPT_KW = dict(N=60, a_max_g=20.0, w_range=1.0, w_effort=1e-3,
              w_time=1e-4, angle_tol_deg=5.0, max_iter=300, verbose=True)

# ============================================================
# Figure 1 — single trajectory optimisation
# ============================================================
print('\n===== Figure 1 =====')
p1 = Params(**{**BASE.__dict__,
               'hit_gamma_v': -45.0,
               'hit_gamma_h': 45.0})
r1 = optimize_trajectory(p1, **OPT_KW)
s1 = r1['summary']
print(f"  miss={s1['optimised_miss_m']} m  RoL={s1['range_on_line_km']} km  "
      f"solved={s1['solved']}  gV_err={s1.get('gamma_v_err_deg')}°")

plot_optimised(r1, title='Direct-collocation — (20 km N, 20 km E), γV=−45°, γH=45°')
plt.savefig('_fig1.png', dpi=150, bbox_inches='tight')
plt.close('all')
print('  Figure 1 saved.')

FIG1_CAP = (
    f"Figure 1 — Direct-collocation trajectory optimiser (v3.0): "
    f"(20 km N, 20 km E) target, γV=−45°, γH=45°. "
    f"Nominal IACPN miss: {s1['nominal_miss_m']:.0f} m → "
    f"optimised miss: {s1['optimised_miss_m']:.2f} m  "
    f"(solved: {s1['solved']}). "
    f"RangeOnLine: {s1['range_on_line_km']:.2f} km. "
    f"Flight time: {s1['optimised_flight_time_s']:.1f} s."
)

# ============================================================
# Figure 3 — optimizer grid, angle-constrained
# ============================================================
print('\n===== Figure 3 =====')
from dataclasses import replace as _rep
base_grid = _rep(BASE, x_target=0, y_target=0)  # will be overridden per run

grid_points = [
    (20, 20, 0),
    ( 0, 25, 0),
    (25,  5, 0),
]
grid_angles = [
    (-45, 45.0),
    (-45, None),
]

runs3 = run_grid(
    _rep(BASE, x_target=0, y_target=0, hit_gamma_v=None, hit_gamma_h=None),
    grid_points,
    hit_angles=grid_angles,
    optimise=True,
    opt_kwargs={**OPT_KW, 'verbose': False},
)

plot_grid(runs3, title='Optimizer grid (v3.0) — 3 targets × 2 angle cases')
plt.savefig('_fig3.png', dpi=150, bbox_inches='tight')
plt.close('all')
print('  Figure 3 saved.')

n_solved3 = sum(1 for r in runs3 if r['summary'].get('solved'))
FIG3_CAP = (
    f"Figure 3 — Optimizer grid (v3.0, direct collocation, RangeOnLine merit). "
    f"Three target points: (20,20,0), (0,25,0), (25,5,0) km × two angle cases "
    f"(γV=−45° with γH=45°; γV=−45° unconstrained). "
    f"{n_solved3}/{len(runs3)} runs solved (miss < 1 m). "
    f"Green bars = solved, red = not solved."
)

# ============================================================
# Figure 4 — optimizer grid, tighter tolerance
# ============================================================
print('\n===== Figure 4 =====')
runs4 = run_grid(
    _rep(BASE, x_target=0, y_target=0, hit_gamma_v=None, hit_gamma_h=None),
    grid_points,
    hit_angles=grid_angles,
    optimise=True,
    opt_kwargs={**OPT_KW, 'angle_tol_deg': 3.0, 'verbose': False},
)

plot_grid(runs4, title='Optimizer grid (v3.0) — tighter angle_tol_deg=3°')
plt.savefig('_fig4.png', dpi=150, bbox_inches='tight')
plt.close('all')
print('  Figure 4 saved.')

n_solved4 = sum(1 for r in runs4 if r['summary'].get('solved'))
FIG4_CAP = (
    f"Figure 4 — Optimizer grid (v3.0), tighter tolerance angle_tol_deg=3°. "
    f"Same targets and angle cases as Figure 3. "
    f"{n_solved4}/{len(runs4)} runs solved. "
    f"ΔγV labels show residual impact angle error."
)

# ============================================================
# Replace images and update captions in docx
# ============================================================
print('\n===== Updating docx =====')
import docx as _docx
from docx.oxml.ns import qn

DOCX = r'files\Booster_Impact_Simulation_Documentation.docx'
doc = _docx.Document(DOCX)
paras = doc.paragraphs

def replace_image(doc, para, img_path):
    """Replace inline image in paragraph by swapping the blob in-place."""
    blips = para._element.findall('.//' + qn('a:blip'))
    if not blips:
        print(f'  WARNING: no blip found in paragraph')
        return False
    r_embed = ('{http://schemas.openxmlformats.org/officeDocument/2006/'
               'relationships}embed')
    rId = blips[0].get(r_embed)
    img_part = doc.part.related_parts[rId]
    with open(img_path, 'rb') as f:
        img_part._blob = f.read()
    return True

def set_text(p, text):
    for run in p.runs:
        run.text = ''
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)

# Replace Fig 1 (para 202), Fig 3 (para 225), Fig 4 (para 227)
imgs = [
    (202, '_fig1.png', 203, FIG1_CAP),
    (225, '_fig3.png', 226, FIG3_CAP),
    (227, '_fig4.png', 228, FIG4_CAP),
]

for img_para_idx, img_file, cap_idx, cap_text in imgs:
    ok = replace_image(doc, paras[img_para_idx], img_file)
    print(f'  Para {img_para_idx}: image replaced = {ok}')
    set_text(paras[cap_idx], cap_text)
    print(f'  Para {cap_idx}: caption updated')

doc.save(DOCX)
print('  docx saved.')

# Clean up temp PNG files
for f in ['_fig1.png', '_fig3.png', '_fig4.png']:
    if os.path.exists(f):
        os.remove(f)
print('  Temp PNGs removed.')
print('\nDone.')
